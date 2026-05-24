import io
from pathlib import Path

import pandas as pd
from pypdf import PdfReader
from docx import Document


class UnsupportedFileTypeError(Exception):
    pass


def extract_text_from_bytes(filename: str, content: bytes) -> str:
    """
    根据文件后缀，将不同格式的文件统一解析成纯文本。
    支持：
    - .txt
    - .md
    - .pdf
    - .docx
    - .csv
    - .xlsx
    - .xls
    """

    suffix = Path(filename).suffix.lower()

    if suffix in [".txt", ".md"]:
        return _extract_text_file(content)

    if suffix == ".pdf":
        return _extract_pdf(content)

    if suffix == ".docx":
        return _extract_docx(content)

    if suffix == ".csv":
        return _extract_csv(content)

    if suffix in [".xlsx", ".xls"]:
        return _extract_excel(content, suffix)

    if suffix == ".doc":
        raise UnsupportedFileTypeError(
            "暂不支持旧版 .doc 文件，请先另存为 .docx 后再上传。"
        )

    raise UnsupportedFileTypeError(
        f"暂不支持该文件类型：{suffix}。当前支持 txt、md、pdf、docx、csv、xlsx、xls。"
    )


def _extract_text_file(content: bytes) -> str:
    """
    解析 txt / md。
    做几个常见编码兜底，避免中文文件乱码。
    """

    for encoding in ["utf-8-sig", "utf-8", "gbk", "gb18030"]:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    return content.decode("utf-8", errors="ignore")


def _extract_pdf(content: bytes) -> str:
    """
    解析 PDF。
    注意：如果 PDF 是扫描图片型 PDF，这里可能提取不到文字。
    """

    reader = PdfReader(io.BytesIO(content))

    pages = []

    for index, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""

        if page_text.strip():
            pages.append(f"【第 {index + 1} 页】\n{page_text}")

    return "\n\n".join(pages)


def _extract_docx(content: bytes) -> str:
    """
    解析 docx。
    提取正文段落和表格内容。
    """

    document = Document(io.BytesIO(content))

    parts = []

    # 1. 提取普通段落
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # 2. 提取表格
    for table_index, table in enumerate(document.tables):
        parts.append(f"\n【Word 表格 {table_index + 1}】")

        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                row_text.append(cell_text)

            if any(row_text):
                parts.append(" | ".join(row_text))

    return "\n".join(parts)


def _extract_csv(content: bytes) -> str:
    """
    解析 csv 表格。
    先尝试不同编码，再用 pandas 读取。
    """

    last_error = None

    for encoding in ["utf-8-sig", "utf-8", "gbk", "gb18030"]:
        try:
            text = content.decode(encoding)
            df = pd.read_csv(io.StringIO(text))
            return _dataframe_to_text(df, title="CSV 表格")
        except Exception as e:
            last_error = e

    raise ValueError(f"CSV 文件解析失败：{last_error}")


def _extract_excel(content: bytes, suffix: str) -> str:
    """
    解析 Excel。
    xlsx 使用 openpyxl，xls 使用 xlrd。
    """

    engine = "openpyxl" if suffix == ".xlsx" else "xlrd"

    sheets = pd.read_excel(
        io.BytesIO(content),
        sheet_name=None,
        engine=engine,
    )

    parts = []

    for sheet_name, df in sheets.items():
        parts.append(_dataframe_to_text(df, title=f"Excel 工作表：{sheet_name}"))

    return "\n\n".join(parts)


def _dataframe_to_text(df: pd.DataFrame, title: str = "表格") -> str:
    """
    把 DataFrame 转成适合 embedding 的文本。
    """

    if df.empty:
        return f"【{title}】\n空表格"

    # 避免表格太大导致文本过长，先限制前 200 行
    df = df.head(200)

    # 把 NaN 转成空字符串
    df = df.fillna("")

    lines = [f"【{title}】"]

    # 表头
    columns = [str(col) for col in df.columns]
    lines.append(" | ".join(columns))

    # 每一行
    for _, row in df.iterrows():
        values = [str(value) for value in row.tolist()]
        lines.append(" | ".join(values))

    return "\n".join(lines)